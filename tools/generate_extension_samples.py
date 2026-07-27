"""8교시 확장 문서 사진·DOCX·PDF·정답 JSON을 재생성한다."""

from __future__ import annotations

import json
import random
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFilter, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.lib.styles import ParagraphStyle

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from src.document_examples import DOCUMENT_EXAMPLES


OUT = ROOT / "sample_docs" / "extensions"
FORMAT_OUT = ROOT / "sample_docs" / "formats"
JSON_OUT = ROOT / "sample_outputs" / "extensions"
INK = "#173B57"
TEAL = "#0F766E"
LIGHT = "#E8F0F4"
MUTED = "#667085"


def find_korean_font(*, bold: bool = False) -> str:
    candidates = (
        [
            "/System/Library/Fonts/AppleSDGothicNeo.ttc",
            "/System/Library/Fonts/Supplemental/AppleGothic.ttf",
        ]
        if not bold
        else [
            "/System/Library/Fonts/AppleSDGothicNeo.ttc",
            "/System/Library/Fonts/Supplemental/AppleGothic.ttf",
        ]
    )
    for candidate in candidates:
        if Path(candidate).is_file():
            return candidate
    raise FileNotFoundError("한글 폰트가 필요합니다.")


def font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(find_korean_font(bold=bold), size)


def money(value: int) -> str:
    return f"{value:,}원"


def document_canvas(title: str, code: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (1180, 1560), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 1180, 88), fill=INK)
    draw.text((58, 27), "교육용 합성 문서 · 개인정보 없음", font=font(25), fill="white")
    draw.text((58, 142), title, font=font(58, bold=True), fill=INK)
    draw.text((825, 157), code, font=font(25), fill=MUTED)
    draw.line((58, 230, 1122, 230), fill=TEAL, width=5)
    return image, draw


def draw_table(
    draw: ImageDraw.ImageDraw,
    *,
    top: int,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
) -> int:
    left = 58
    row_height = 72
    x_positions = [left]
    for width in widths:
        x_positions.append(x_positions[-1] + width)
    draw.rectangle((left, top, x_positions[-1], top + row_height), fill=INK)
    for index, text in enumerate(headers):
        draw.text(
            (x_positions[index] + 12, top + 19),
            text,
            font=font(25, bold=True),
            fill="white",
        )
    y = top + row_height
    for row_index, row in enumerate(rows):
        fill = "#F7FAFC" if row_index % 2 == 0 else "white"
        draw.rectangle(
            (left, y, x_positions[-1], y + row_height),
            fill=fill,
            outline="#C8D2DC",
            width=2,
        )
        for index, text in enumerate(row):
            draw.text(
                (x_positions[index] + 12, y + 20),
                text,
                font=font(24),
                fill="#1D2939",
            )
            draw.line(
                (x_positions[index], y, x_positions[index], y + row_height),
                fill="#C8D2DC",
                width=2,
            )
        y += row_height
    return y


def quotation_page() -> Image.Image:
    data = DOCUMENT_EXAMPLES["quotation"]
    image, draw = document_canvas("견 적 서", data["document_number"])
    labels = [
        ("공급자", data["supplier"]),
        ("수신", data["customer"]),
        ("견적일", data["quote_date"]),
        ("유효기간", data["valid_until"]),
    ]
    y = 282
    for label, value in labels:
        draw.text((62, y), label, font=font(25, bold=True), fill=MUTED)
        draw.text((230, y), value, font=font(29), fill="#1D2939")
        y += 55
    rows = [
        [
            item["name"],
            str(item["quantity"]),
            money(item["unit_price"]),
            money(item["line_total"]),
        ]
        for item in data["items"]
    ]
    y = draw_table(
        draw,
        top=535,
        headers=["품목", "수량", "단가", "금액"],
        rows=rows,
        widths=[460, 120, 230, 250],
    )
    totals = [
        ("공급가액", data["subtotal"]),
        ("부가세", data["tax"]),
        ("합계", data["total_amount"]),
    ]
    y += 45
    for label, value in totals:
        draw.text((690, y), label, font=font(27, bold=True), fill=INK)
        draw.text((900, y), money(value), font=font(29, bold=True), fill=TEAL)
        y += 58
    draw.rounded_rectangle((58, 1210, 1122, 1405), radius=18, fill=LIGHT)
    draw.text((90, 1245), "특이사항", font=font(28, bold=True), fill=INK)
    draw.text(
        (90, 1300),
        "납품 장소와 설치 일정은 발주 전 협의합니다.",
        font=font(26),
        fill="#344054",
    )
    return image


def application_page() -> Image.Image:
    data = DOCUMENT_EXAMPLES["application"]
    image, draw = document_canvas("교육 참가 신청서", data["application_id"])
    rows = [
        ("신청자", data["applicant_name"]),
        ("소속", data["department"]),
        ("신청 과정", data["requested_program"]),
        ("희망 일자", data["requested_date"]),
        ("관리자 승인", data["manager_approval"]),
    ]
    y = 300
    for index, (label, value) in enumerate(rows):
        height = 104 if index != 2 else 130
        draw.rectangle((58, y, 330, y + height), fill=LIGHT, outline="#B8C5D1", width=2)
        draw.rectangle((330, y, 1122, y + height), fill="white", outline="#B8C5D1", width=2)
        draw.text((88, y + 34), label, font=font(27, bold=True), fill=INK)
        draw.text((370, y + 34), str(value), font=font(29), fill="#1D2939")
        y += height
    draw.rounded_rectangle((58, 1000, 1122, 1300), radius=16, outline=INK, width=3)
    draw.text((90, 1040), "개인정보 수집·이용 동의", font=font(30, bold=True), fill=INK)
    draw.text(
        (90, 1100),
        "교육 운영과 출석 확인에 필요한 최소 정보만 사용합니다.",
        font=font(25),
        fill="#344054",
    )
    checked = "☑" if data["privacy_consent"] else "☐"
    draw.text((90, 1205), f"{checked} 동의함", font=font(31, bold=True), fill=TEAL)
    draw.text((820, 1370), "신청자 서명  김하늘", font=font(27), fill="#344054")
    return image


def transaction_page() -> Image.Image:
    data = DOCUMENT_EXAMPLES["transaction_statement"]
    image, draw = document_canvas("거 래 명 세 서", data["statement_number"])
    labels = [
        ("공급자", data["supplier"]),
        ("공급받는 자", data["customer"]),
        ("거래일", data["transaction_date"]),
    ]
    y = 295
    for label, value in labels:
        draw.text((62, y), label, font=font(25, bold=True), fill=MUTED)
        draw.text((260, y), value, font=font(29), fill="#1D2939")
        y += 58
    rows = [
        [
            item["name"],
            str(item["quantity"]),
            money(item["unit_price"]),
            money(item["line_total"]),
        ]
        for item in data["items"]
    ]
    y = draw_table(
        draw,
        top=515,
        headers=["품목", "수량", "단가", "공급가액"],
        rows=rows,
        widths=[460, 120, 230, 250],
    )
    y += 55
    for label, value in (
        ("공급가액", data["subtotal"]),
        ("세액", data["tax"]),
        ("합계금액", data["total_amount"]),
    ):
        draw.text((680, y), label, font=font(28, bold=True), fill=INK)
        draw.text((900, y), money(value), font=font(30, bold=True), fill=TEAL)
        y += 64
    draw.text((58, 1400), "인수자 확인  ____________________", font=font(26), fill="#344054")
    return image


def photographed(page: Image.Image, *, seed: int) -> Image.Image:
    """합성 원본을 책상 위에서 촬영한 한 장짜리 사진처럼 만든다."""

    random.seed(seed)
    desk = Image.new("RGB", (1500, 1800), "#A87C52")
    pixels = desk.load()
    for y in range(desk.height):
        band = int(9 * (y / 130)) % 2
        for x in range(desk.width):
            noise = random.randint(-5, 5)
            base = (168 + noise - band * 4, 124 + noise, 82 + noise)
            pixels[x, y] = base
    page = page.resize((1040, 1375), Image.Resampling.LANCZOS)
    page = page.rotate(-2.3 + seed * 0.4, expand=True, fillcolor="#F7F3ED")
    shadow = Image.new("RGBA", page.size, (0, 0, 0, 0))
    ImageDraw.Draw(shadow).rounded_rectangle(
        (18, 18, page.width - 8, page.height - 8),
        radius=12,
        fill=(0, 0, 0, 95),
    )
    shadow = shadow.filter(ImageFilter.GaussianBlur(18))
    x = (desk.width - page.width) // 2
    y = 170
    desk.paste(shadow, (x + 16, y + 22), shadow)
    desk.paste(page, (x, y))
    return desk


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_width(cell, dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.first_child_found_in("w:tcW")
    if width is None:
        width = OxmlElement("w:tcW")
        tc_pr.append(width)
    width.set(qn("w:w"), str(dxa))
    width.set(qn("w:type"), "dxa")


def set_run(run, *, size: float = 11, bold: bool = False, color: str = "1D2939"):
    run.font.name = "Arial Unicode MS"
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), "Arial Unicode MS")
    fonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    fonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def create_application_docx(path: Path) -> None:
    """이미지 기반 Word 문서 예시를 만든다.

    같은 ``.docx``라도 본문이 편집 가능한 텍스트가 아니라 한 장의 이미지일
    수 있다는 점을 체험하도록 설계한 샘플이다.
    """

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("WORD SAMPLE · IMAGE-BASED FORM")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("667085")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Synthetic training document · no personal data")
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string("667085")

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_picture(str(OUT / "application_form_flat.png"), width=Inches(6.95))
    drawing = run._element.xpath(".//wp:docPr")
    if drawing:
        drawing[0].set("descr", "교육용 합성 교육 참가 신청서")
    doc.save(path)


def create_transaction_pdf(path: Path) -> None:
    data = DOCUMENT_EXAMPLES["transaction_statement"]
    font_path = "/System/Library/Fonts/Supplemental/AppleGothic.ttf"
    if not Path(font_path).is_file():
        font_path = find_korean_font()
    pdfmetrics.registerFont(TTFont("Korean", font_path))
    styles = {
        "title": ParagraphStyle(
            "title",
            fontName="Korean",
            fontSize=24,
            leading=30,
            textColor=colors.HexColor(INK),
            alignment=1,
        ),
        "body": ParagraphStyle(
            "body",
            fontName="Korean",
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#1D2939"),
        ),
    }
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=48,
        rightMargin=48,
        topMargin=42,
        bottomMargin=42,
        title="교육용 거래명세서",
        author="Document AI 교육",
    )
    story = [
        Paragraph("거 래 명 세 서", styles["title"]),
        Spacer(1, 14),
        Paragraph(
            f"문서번호: {data['statement_number']}　 거래일: {data['transaction_date']}",
            styles["body"],
        ),
        Paragraph(
            f"공급자: {data['supplier']}　 공급받는 자: {data['customer']}",
            styles["body"],
        ),
        Spacer(1, 16),
    ]
    rows = [["품목", "수량", "단가", "공급가액"]]
    rows.extend(
        [
            item["name"],
            str(item["quantity"]),
            money(item["unit_price"]),
            money(item["line_total"]),
        ]
        for item in data["items"]
    )
    rows.extend(
        [
            ["", "", "공급가액", money(data["subtotal"])],
            ["", "", "세액", money(data["tax"])],
            ["", "", "합계금액", money(data["total_amount"])],
        ]
    )
    table = Table(rows, colWidths=[230, 60, 105, 110], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "Korean"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(INK)),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B8C5D1")),
                ("ROWBACKGROUNDS", (0, 1), (-1, 2), [colors.white, colors.HexColor("#F7FAFC")]),
                ("TOPPADDING", (0, 0), (-1, -1), 9),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
            ]
        )
    )
    story.extend(
        [
            table,
            Spacer(1, 22),
            Paragraph("교육용 합성 문서 · 실제 개인정보 없음", styles["body"]),
        ]
    )
    doc.build(story)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    FORMAT_OUT.mkdir(parents=True, exist_ok=True)
    JSON_OUT.mkdir(parents=True, exist_ok=True)

    pages = {
        "quotation": quotation_page(),
        "application_form": application_page(),
        "transaction_statement": transaction_page(),
    }
    for index, (name, page) in enumerate(pages.items(), start=1):
        page.save(OUT / f"{name}_flat.png", optimize=True)
        photographed(page, seed=index).save(OUT / f"{name}_photo.png", optimize=True)

    create_application_docx(FORMAT_OUT / "application_form.docx")
    create_transaction_pdf(FORMAT_OUT / "transaction_statement.pdf")

    for document_type, data in DOCUMENT_EXAMPLES.items():
        payload = {
            **data,
            "_provenance": {
                "fixture_type": "synthetic_training_document",
                "created_at": "2026-07-28",
                "contains_personal_data": False,
                "reviewer": "course maintainer",
            },
        }
        (JSON_OUT / f"{document_type}.json").write_text(
            json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )

    print("확장 문서 사진·DOCX·PDF·정답 JSON을 생성했습니다.")


if __name__ == "__main__":
    main()
